"""
DOCX文档解析器
提取Word文档内容和格式信息
"""

import uuid
import logging
from typing import List, Optional, Dict, Any
from pathlib import Path

try:
    from docx import Document
    HAS_PYTHON_DOCX = True
except ImportError:
    HAS_PYTHON_DOCX = False
    logging.warning("python-docx not installed, DOCX parsing disabled")

from ..models.document_models import DocumentParseResult, PageInfo, TextBlock, DocumentType


logger = logging.getLogger(__name__)


class DOCXParser:
    """DOCX文档解析器"""
    
    def parse(self, file_path: str, document_id: Optional[str] = None) -> DocumentParseResult:
        """
        解析DOCX文档
        
        Args:
            file_path: DOCX文件路径
            document_id: 文档ID，如果为None则自动生成
            
        Returns:
            DocumentParseResult: 解析结果
        """
        if not HAS_PYTHON_DOCX:
            raise ImportError("需要安装 python-docx 来解析DOCX文档")
            
        if document_id is None:
            document_id = str(uuid.uuid4())
            
        logger.info(f"开始解析DOCX文档: {file_path}")
        
        import time
        start_time = time.time()
        
        doc = Document(file_path)
        
        # DOCX没有页面概念，我们按段落分块，然后模拟分页
        all_text_blocks = []
        current_char_offset = 0
        
        # 处理所有段落
        for para_idx, paragraph in enumerate(doc.paragraphs):
            if paragraph.text.strip():
                block_id = f"{document_id}_para{para_idx}"
                
                # 创建文本块
                block = TextBlock(
                    text=paragraph.text.strip(),
                    page_number=1,  # DOCX暂时都放在第1页
                    block_id=block_id,
                    char_start=current_char_offset,
                    char_end=current_char_offset + len(paragraph.text),
                    confidence=1.0,
                    block_type="paragraph",
                    is_ocr=False
                )
                
                all_text_blocks.append(block)
                current_char_offset += len(paragraph.text)
        
        # 处理表格
        for table_idx, table in enumerate(doc.tables):
            for row_idx, row in enumerate(table.rows):
                for cell_idx, cell in enumerate(row.cells):
                    if cell.text.strip():
                        block_id = f"{document_id}_table{table_idx}_row{row_idx}_cell{cell_idx}"
                        
                        block = TextBlock(
                            text=cell.text.strip(),
                            page_number=1,
                            block_id=block_id,
                            char_start=current_char_offset,
                            char_end=current_char_offset + len(cell.text),
                            confidence=1.0,
                            block_type="table_cell",
                            is_ocr=False
                        )
                        
                        all_text_blocks.append(block)
                        current_char_offset += len(cell.text)
        
        # 创建页面信息（DOCX模拟为单页）
        page_info = PageInfo(
            page_number=1,
            width=595,  # A4宽度
            height=842, # A4高度
            text_blocks=all_text_blocks,
            has_images=False,  # 暂不处理图片
            image_count=0
        )
        
        total_text_length = sum(len(block.text) for block in all_text_blocks)
        processing_time = time.time() - start_time
        
        result = DocumentParseResult(
            document_id=document_id,
            document_type=DocumentType.DOCX,
            total_pages=1,
            total_text_length=total_text_length,
            pages=[page_info],
            all_text_blocks=all_text_blocks,
            metadata={
                "file_path": file_path,
                "file_size": Path(file_path).stat().st_size,
                "parser": "python-docx",
                "paragraphs_count": len(doc.paragraphs),
                "tables_count": len(doc.tables)
            },
            processing_time=processing_time
        )
        
        logger.info(f"DOCX解析完成，{len(doc.paragraphs)} 个段落，{len(doc.tables)} 个表格，{total_text_length} 字符，耗时 {processing_time:.2f}秒")
        return result